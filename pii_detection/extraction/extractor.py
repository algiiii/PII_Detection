"""Minimal text extraction from real document formats (block B3).

Reads a born-digital file — PDF, Word (``.docx`` and legacy ``.doc``),
spreadsheets (``.xlsx``/``.xlsm``/``.ods``) or plain text — into the
:class:`~pii_detection.detection.types.NormalizedDocument` the detection layer
(B4) consumes. This is the real "receive a document, unpack and read it"
capability, kept deliberately minimal: **no OCR** for scanned files and **no
layout/table reconstruction** (later concerns, e.g. via Docling). The heavy
readers are imported lazily, so importing this module — and building its docs —
needs no extraction dependency; only actually reading a file does. Legacy ``.doc``
relies on the ``antiword`` system binary, present in the container image.
"""

from __future__ import annotations

import re
from collections.abc import Callable
from pathlib import Path

from pii_detection.detection.types import NormalizedDocument


class UnsupportedFormatError(ValueError):
    """Raised when a file's extension has no registered extractor."""


_TRAILING_WS = re.compile(r"[ \t]+\n")
_BLANK_RUNS = re.compile(r"\n{3,}")


def normalize_text(raw: str) -> str:
    """Apply a light, predictable normalization to extracted text.

    Unifies newlines, strips trailing spaces on each line and collapses runs of
    blank lines. Deliberately conservative: it only tidies whitespace, it never
    merges or moves tokens, so the text stays faithful to what was in the file.

    :param raw: text as returned by a format reader.
    :returns: the normalized text.
    """
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    text = _TRAILING_WS.sub("\n", text)
    text = _BLANK_RUNS.sub("\n\n", text)
    return text.strip()


def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF with PyMuPDF, in page order (no OCR)."""
    import fitz  # lazy: heavy dependency, only needed for PDFs

    with fitz.open(path) as document:
        return "\n".join(page.get_text() for page in document)


def _extract_docx(path: Path) -> str:
    """Extract paragraph text from a Word ``.docx`` (tables not covered)."""
    from docx import Document  # lazy

    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _extract_txt(path: Path) -> str:
    """Read a plain-text file as UTF-8."""
    return path.read_text(encoding="utf-8")


def _extract_spreadsheet(path: Path) -> str:
    """Flatten a spreadsheet (``.xlsx``/``.xlsm``/``.ods``) to text, sheet by sheet.

    Reuses the ROPA spreadsheet reader (openpyxl/odfpy behind one shape); cells are
    joined by tabs, rows by newlines and sheets by a blank line, so the detection
    layer sees the cell contents as plain text.
    """
    from pii_detection.ropa.ingestion.sheet_reader import read_sheet, sheet_names  # lazy

    sheets = [
        "\n".join("\t".join(row) for row in read_sheet(path, name))
        for name in sheet_names(path)
    ]
    return "\n\n".join(sheets)


def _extract_doc(path: Path) -> str:
    """Extract text from a legacy Word ``.doc`` via the ``antiword`` binary.

    :raises UnsupportedFormatError: if ``antiword`` is not installed (it ships in
        the container image; a local run without it lets the batch skip the file).
    """
    import shutil  # lazy
    import subprocess  # lazy

    if shutil.which("antiword") is None:
        raise UnsupportedFormatError(
            "reading '.doc' requires the 'antiword' binary (present in the container image)"
        )
    result = subprocess.run(
        ["antiword", str(path)], capture_output=True, text=True, check=True
    )
    return result.stdout


_EXTRACTORS: dict[str, Callable[[Path], str]] = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_txt,
    ".xlsx": _extract_spreadsheet,
    ".xlsm": _extract_spreadsheet,
    ".ods": _extract_spreadsheet,
    ".doc": _extract_doc,
}


def supported_suffixes() -> frozenset[str]:
    """:returns: the file extensions the extractor can read."""
    return frozenset(_EXTRACTORS)


def extract_document(path: str | Path) -> NormalizedDocument:
    """Read a document into a :class:`NormalizedDocument` for the detection layer.

    Dispatches on the file extension; the ``document_id`` is the file stem, so a
    rendered corpus file keeps the identity of its source document.

    :param path: path to a supported file (``.pdf``, ``.docx``, ``.doc``,
        ``.xlsx``/``.xlsm``/``.ods`` or ``.txt``).
    :returns: the normalized document (``document_id`` + normalized ``text``).
    :raises FileNotFoundError: if the file does not exist.
    :raises UnsupportedFormatError: if the extension has no registered extractor.
    """
    path = Path(path)
    if not path.is_file():
        raise FileNotFoundError(f"document not found: {path}")
    extractor = _EXTRACTORS.get(path.suffix.lower())
    if extractor is None:
        supported = ", ".join(sorted(_EXTRACTORS))
        raise UnsupportedFormatError(
            f"no extractor for {path.suffix!r}; supported: {supported}"
        )
    return NormalizedDocument(document_id=path.stem, text=normalize_text(extractor(path)))


__all__ = [
    "UnsupportedFormatError",
    "normalize_text",
    "supported_suffixes",
    "extract_document",
]
